"""Construcción de datos de reporte y generación PDF (personalizable, UTF-8, tablas legibles)."""

from __future__ import annotations

import base64
import io
import os
import re
from collections import Counter
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Tuple
from xml.sax.saxutils import escape

import reportlab
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from services.api import MacuinApi

_FONT_BODY = "Helvetica"
_FONT_TITLE = "Helvetica-Bold"


def _register_unicode_font() -> Tuple[str, str]:
    """Registra DejaVu Sans si existe en el paquete reportlab (acentos y ñ en PDF)."""
    global _FONT_BODY, _FONT_TITLE
    try:
        pkg = os.path.dirname(reportlab.__file__)
        regular = os.path.join(pkg, "fonts", "DejaVuSans.ttf")
        bold = os.path.join(pkg, "fonts", "DejaVuSans-Bold.ttf")
        if os.path.isfile(regular):
            pdfmetrics.registerFont(TTFont("DejaVuSans", regular))
            _FONT_BODY = "DejaVuSans"
            if os.path.isfile(bold):
                pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", bold))
                _FONT_TITLE = "DejaVuSans-Bold"
            else:
                _FONT_TITLE = "DejaVuSans"
    except Exception:
        pass
    return _FONT_BODY, _FONT_TITLE


_register_unicode_font()

# Identidad visual MACUIN (sidebar / panel)
MACUIN_NAVY = "#1e3a5f"
MACUIN_NAVY_DARK = "#0f172a"
MACUIN_GOLD = "#fbbf24"
MACUIN_GOLD_DEEP = "#ca8a04"
MACUIN_SLATE = "#334155"
MACUIN_MUTED = "#64748b"
MACUIN_BG = "#f8fafc"

# Cuatro tipos de reporte (rúbrica): id → secciones incluidas en build_report_context
REPORT_PRESETS: Dict[str, Dict[str, Any]] = {
    "pedidos": {
        "label": "Pedidos",
        "descripcion_corta": "Listado de pedidos con cliente, destino y estatus.",
        "sections": ["portada", "resumen", "pedidos"],
        "roles": {"Administrador", "Ventas", "Logística", "Almacén"},
    },
    "catalogo_autopartes": {
        "label": "Catálogo de autopartes",
        "descripcion_corta": "SKU, nombre, precio, categoría y marca.",
        "sections": ["portada", "resumen", "autopartes"],
        "roles": {"Administrador", "Ventas", "Logística", "Almacén"},
    },
    "usuarios_internos": {
        "label": "Usuarios internos",
        "descripcion_corta": "Personal del sistema y roles (solo administración).",
        "sections": ["portada", "resumen", "usuarios"],
        "roles": {"Administrador"},
    },
    "inventario_almacen": {
        "label": "Inventario y almacén",
        "descripcion_corta": "Stock por autoparte y ubicación física.",
        "sections": ["portada", "resumen", "inventarios", "ubicaciones"],
        "roles": {"Administrador", "Almacén"},
    },
}


def report_presets_for_role(rol: str) -> List[Dict[str, Any]]:
    """Lista de presets visibles para el rol (para la vista de reportes)."""
    out: List[Dict[str, Any]] = []
    for pid, meta in REPORT_PRESETS.items():
        allowed = rol in meta["roles"]
        out.append(
            {
                "id": pid,
                "label": meta["label"],
                "descripcion_corta": meta["descripcion_corta"],
                "permitido": allowed,
            }
        )
    return out


def preset_sections_if_allowed(rol: str, preset_id: str) -> Optional[List[str]]:
    meta = REPORT_PRESETS.get(preset_id)
    if not meta or rol not in meta["roles"]:
        return None
    return list(meta["sections"])


def _normalize_section_order(raw: List[str]) -> List[str]:
    """Portada y resumen primero; el resto conserva el orden del formulario."""
    seen: set[str] = set()
    out: List[str] = []
    for x in ("portada", "resumen"):
        if x in raw and x not in seen:
            out.append(x)
            seen.add(x)
    for x in raw:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out


def _build_macuin_bar_chart_png(
    labels: List[str],
    values: List[int],
    title: str,
    subtitle: str = "",
    horizontal: bool = False,
) -> bytes:
    """Gráfica de barras con paleta corporativa (PNG)."""
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    labels = [_safe(x)[:42] for x in labels]
    values = [int(v) for v in values]
    if not labels or not values or len(labels) != len(values):
        labels, values = ["Sin datos"], [0]

    fig, ax = plt.subplots(figsize=(8.4, 4.35), dpi=115)
    fig.patch.set_facecolor(MACUIN_BG)
    ax.set_facecolor("#ffffff")

    if horizontal:
        y_pos = range(len(labels))
        ax.barh(
            list(y_pos),
            values,
            color=MACUIN_GOLD,
            edgecolor=MACUIN_GOLD_DEEP,
            linewidth=0.6,
            height=0.65,
        )
        ax.set_yticks(list(y_pos))
        ax.set_yticklabels(labels, fontsize=9, color=MACUIN_SLATE)
        ax.invert_yaxis()
        ax.set_xlabel("Cantidad", fontsize=9, color=MACUIN_SLATE)
    else:
        x = range(len(labels))
        ax.bar(
            list(x),
            values,
            color=MACUIN_GOLD,
            edgecolor=MACUIN_GOLD_DEEP,
            linewidth=0.6,
            width=0.72,
        )
        ax.set_xticks(list(x))
        ax.set_xticklabels(labels, rotation=28, ha="right", fontsize=9, color=MACUIN_SLATE)
        ax.set_ylabel("Cantidad", fontsize=9, color=MACUIN_SLATE)

    ax.set_title(title, fontsize=13, fontweight="bold", color=MACUIN_NAVY_DARK, pad=12)
    if subtitle:
        fig.text(0.5, 0.02, subtitle, ha="center", fontsize=8.5, color=MACUIN_MUTED)
    ax.grid(axis="y" if not horizontal else "x", alpha=0.28, linestyle="--", color=MACUIN_NAVY)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for s in ("bottom", "left"):
        ax.spines[s].set_color(MACUIN_NAVY)
    ax.tick_params(colors=MACUIN_SLATE)
    fig.subplots_adjust(bottom=0.22 if not horizontal else 0.12, left=0.28 if horizontal else 0.1, right=0.96, top=0.88)

    buf = io.BytesIO()
    fig.savefig(buf, format="png", facecolor=fig.get_facecolor(), edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _chart_sections_for_context(api: MacuinApi) -> List[Dict[str, Any]]:
    """Gráficas de negocio para incrustar en PDF / Word / Excel."""
    out: List[Dict[str, Any]] = []

    try:
        data = api.get("/v1/movimientos-inventario/resumen-por-mes", params={"months": 10})
        labels = list(data.get("labels") or [])
        vals_raw = data.get("values") or []
        vals = [int(v) for v in vals_raw]
        if labels and vals and len(labels) == len(vals):
            png = _build_macuin_bar_chart_png(
                labels,
                vals,
                "Movimientos de inventario por mes",
                "Serie histórica de movimientos registrados en almacén",
            )
            out.append(
                {
                    "type": "chart",
                    "title": "Gráfica — movimientos de inventario",
                    "subtitle": "Últimos meses con actividad registrada",
                    "png_bytes": png,
                    "png_b64": base64.standard_b64encode(png).decode("ascii"),
                }
            )
    except Exception:
        pass

    try:
        vista = api.get("/v1/pedidos/admin/vista", params={"limit": 500}) or []
        ct: Counter[str] = Counter()
        for p in vista:
            if isinstance(p, dict) and p.get("estatus"):
                ct[str(p.get("estatus"))] += 1
        if ct:
            labels = list(ct.keys())
            vals = [int(ct[k]) for k in labels]
            horiz = max((len(str(x)) for x in labels), default=0) > 16
            png = _build_macuin_bar_chart_png(
                labels,
                vals,
                "Pedidos por estatus",
                "Distribución de pedidos según estatus operativo",
                horizontal=horiz,
            )
            out.append(
                {
                    "type": "chart",
                    "title": "Gráfica — pedidos por estatus",
                    "subtitle": "Basado en el listado administrativo de pedidos",
                    "png_bytes": png,
                    "png_b64": base64.standard_b64encode(png).decode("ascii"),
                }
            )
    except Exception:
        pass

    return out


def available_sections_for_role(rol: str) -> List[Dict[str, str]]:
    """Metadatos de secciones que cada rol puede incluir en el PDF."""
    base = [
        {"id": "portada", "label": "Portada (título, fecha, usuario)"},
        {"id": "resumen", "label": "Resumen ejecutivo (KPIs texto)"},
    ]
    if rol == "Administrador":
        base += [
            {"id": "usuarios", "label": "Tabla de usuarios internos"},
            {"id": "roles", "label": "Roles del sistema"},
            {"id": "autopartes", "label": "Catálogo de autopartes (primeras N filas)"},
            {"id": "pedidos", "label": "Pedidos (primeras N filas)"},
            {"id": "parametros", "label": "Parámetros de sistema"},
            {"id": "mensajes_contacto", "label": "Mensajes de contacto del portal web"},
        ]
    elif rol == "Ventas":
        base += [
            {"id": "clientes", "label": "Clientes"},
            {"id": "pedidos", "label": "Pedidos"},
            {"id": "autopartes", "label": "Catálogo (consulta)"},
        ]
    elif rol == "Logística":
        base += [
            {"id": "pedidos", "label": "Pedidos y estatus"},
            {"id": "guias", "label": "Guías de envío"},
            {"id": "direcciones", "label": "Direcciones (listado)"},
        ]
    elif rol == "Almacén":
        base += [
            {"id": "inventarios", "label": "Inventario actual"},
            {"id": "ubicaciones", "label": "Ubicaciones"},
            {"id": "pedidos", "label": "Pedidos (surtido/empaque)"},
            {"id": "autopartes", "label": "Autopartes (referencia)"},
        ]
    return base


def _lim_rows(rows: List[List[Any]], max_rows: int) -> List[List[Any]]:
    """Limita filas de datos (sin cabecera)."""
    if len(rows) <= max_rows:
        return rows
    head = rows[:max_rows]
    n_extra = len(rows) - max_rows
    nc = len(rows[0]) if rows else 1
    note = [f"… ({n_extra} filas adicionales no mostradas)"] + [""] * (nc - 1)
    head.append(note[:nc])
    return head


def _safe(s: Any) -> str:
    if s is None:
        return "—"
    return str(s)


def _lookup_maps(api: MacuinApi) -> Dict[str, Dict[int, str]]:
    """Mapas id → texto legible para sustituir FKs en tablas."""
    maps: Dict[str, Dict[int, str]] = {
        "rol": {},
        "categoria": {},
        "marca": {},
        "estatus": {},
        "cliente": {},
        "ubicacion": {},
        "autoparte": {},
        "pedido_folio": {},
    }
    try:
        for r in api.get("/v1/roles/") or []:
            rid = r.get("id") if isinstance(r, dict) else getattr(r, "id", None)
            if rid is not None:
                nom = (r.get("nombre_rol") if isinstance(r, dict) else getattr(r, "nombre_rol", "")) or ""
                maps["rol"][int(rid)] = nom
    except Exception:
        pass
    try:
        for c in api.get("/v1/categorias/") or []:
            cid = c.get("id") if isinstance(c, dict) else getattr(c, "id", None)
            if cid is not None:
                maps["categoria"][int(cid)] = (c.get("nombre") if isinstance(c, dict) else getattr(c, "nombre", "")) or ""
    except Exception:
        pass
    try:
        for m in api.get("/v1/marcas/") or []:
            mid = m.get("id") if isinstance(m, dict) else getattr(m, "id", None)
            if mid is not None:
                maps["marca"][int(mid)] = (m.get("nombre") if isinstance(m, dict) else getattr(m, "nombre", "")) or ""
    except Exception:
        pass
    try:
        for e in api.get("/v1/estatus-pedido/") or []:
            eid = e.get("id") if isinstance(e, dict) else getattr(e, "id", None)
            if eid is not None:
                maps["estatus"][int(eid)] = (e.get("nombre") if isinstance(e, dict) else getattr(e, "nombre", "")) or ""
    except Exception:
        pass
    try:
        for c in api.get("/v1/clientes/") or []:
            cid = c.get("id") if isinstance(c, dict) else getattr(c, "id", None)
            if cid is not None:
                maps["cliente"][int(cid)] = (c.get("nombre") if isinstance(c, dict) else getattr(c, "nombre", "")) or ""
    except Exception:
        pass
    try:
        for u in api.get("/v1/ubicaciones/") or []:
            uid = u.get("id") if isinstance(u, dict) else getattr(u, "id", None)
            if uid is not None:
                p, es, nv = (
                    (u.get("pasillo"), u.get("estante"), u.get("nivel"))
                    if isinstance(u, dict)
                    else (getattr(u, "pasillo", ""), getattr(u, "estante", ""), getattr(u, "nivel", None))
                )
                lbl = f"{p}-{es}" + (f"-{nv}" if nv else "")
                maps["ubicacion"][int(uid)] = lbl
    except Exception:
        pass
    try:
        for a in api.get("/v1/autopartes/") or []:
            aid = a.get("id") if isinstance(a, dict) else getattr(a, "id", None)
            if aid is not None:
                sku = (a.get("sku_codigo") if isinstance(a, dict) else getattr(a, "sku_codigo", "")) or ""
                nom = (a.get("nombre") if isinstance(a, dict) else getattr(a, "nombre", "")) or ""
                maps["autoparte"][int(aid)] = f"{sku} — {nom[:50]}"
    except Exception:
        pass
    try:
        vista = api.get("/v1/pedidos/admin/vista", params={"limit": 500}) or []
        for p in vista:
            if isinstance(p, dict) and p.get("id") is not None:
                maps["pedido_folio"][int(p["id"])] = _safe(p.get("folio"))
    except Exception:
        pass
    return maps


def build_report_context(
    api: MacuinApi,
    rol: str,
    options: Dict[str, Any],
) -> Dict[str, Any]:
    """options: sections (list), max_rows, title, notes, include_charts (bool)."""
    sections: List[Dict[str, Any]] = []
    title = (options.get("title") or f"Reporte Macuin — {rol}").strip()[:200]
    notes = (options.get("notes") or "").strip()[:2000]
    max_rows = int(options.get("max_rows") or 80)
    max_rows = max(5, min(max_rows, 500))
    raw_sections = list(options.get("sections") or ["portada", "resumen"])
    chosen_list = _normalize_section_order(raw_sections)
    include_charts = bool(options.get("include_charts"))
    maps = _lookup_maps(api)

    def add_table(title_p: str, headers: List[str], data_rows: List[List[Any]]):
        data_rows = _lim_rows(data_rows, max_rows)
        sections.append({"type": "table", "title": title_p, "headers": headers, "rows": data_rows})

    for sid in chosen_list:
        if sid == "portada":
            sections.append(
                {
                    "type": "text",
                    "title": "Portada",
                    "body": (
                        f"{title}\n"
                        f"Sistema: MACUIN Autopartes\n"
                        f"Generado: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n"
                        f"Perfil: {rol}"
                    ),
                }
            )
        elif sid == "resumen":
            sections.append(
                {
                    "type": "text",
                    "title": "Notas / resumen ejecutivo",
                    "body": notes or "Sin notas adicionales.",
                }
            )

    if include_charts:
        sections.extend(_chart_sections_for_context(api))

    try:
        for sid in chosen_list:
            if sid in ("portada", "resumen"):
                continue
            if sid == "usuarios" and rol == "Administrador":
                raw = api.get("/v1/usuarios/")
                usuarios = raw.get("usuarios", []) if isinstance(raw, dict) else []
                rows = []
                for u in usuarios:
                    if not isinstance(u, dict):
                        u = {k: getattr(u, k, None) for k in ("id", "nombre", "apellidos", "email", "rol_id", "activo")}
                    rid = u.get("rol_id")
                    rol_txt = maps["rol"].get(int(rid), "—") if rid is not None else "—"
                    rows.append(
                        [
                            _safe(u.get("nombre")),
                            _safe(u.get("apellidos")),
                            _safe(u.get("email")),
                            rol_txt,
                            "Sí" if u.get("activo") else "No",
                        ]
                    )
                add_table("Usuarios internos", ["Nombre", "Apellidos", "Correo", "Rol", "Activo"], rows)

            elif sid == "roles" and rol == "Administrador":
                roles = api.get("/v1/roles/") or []
                rows = []
                for r in roles:
                    if not isinstance(r, dict):
                        r = {k: getattr(r, k, None) for k in ("nombre_rol", "descripcion")}
                    rows.append([_safe(r.get("nombre_rol")), _safe(r.get("descripcion"))])
                add_table("Roles del sistema", ["Rol", "Descripción"], rows)

            elif sid == "clientes" and rol in ("Administrador", "Ventas"):
                clientes = api.get("/v1/clientes/") or []
                rows = []
                for c in clientes:
                    if not isinstance(c, dict):
                        c = {k: getattr(c, k, None) for k in ("nombre", "email", "telefono", "activo", "notas")}
                    rows.append(
                        [
                            _safe(c.get("nombre")),
                            _safe(c.get("email")),
                            _safe(c.get("telefono")),
                            "Sí" if c.get("activo") else "No",
                            (_safe(c.get("notas"))[:80] + "…")
                            if c.get("notas") and len(_safe(c.get("notas"))) > 80
                            else _safe(c.get("notas")),
                        ]
                    )
                add_table("Clientes", ["Nombre", "Correo", "Teléfono", "Activo", "Notas (extracto)"], rows)

            elif sid == "autopartes":
                partes = api.get("/v1/autopartes/") or []
                rows = []
                for p in partes:
                    if not isinstance(p, dict):
                        p = {k: getattr(p, k, None) for k in ("sku_codigo", "nombre", "precio_unitario", "categoria_id", "marca_id")}
                    cid, mid = p.get("categoria_id"), p.get("marca_id")
                    cat = maps["categoria"].get(int(cid), "—") if cid is not None else "—"
                    mar = maps["marca"].get(int(mid), "—") if mid is not None else "—"
                    rows.append(
                        [
                            _safe(p.get("sku_codigo")),
                            _safe(p.get("nombre")),
                            _safe(p.get("precio_unitario")),
                            cat,
                            mar,
                        ]
                    )
                add_table("Catálogo de autopartes", ["SKU", "Nombre", "Precio unit. (MXN)", "Categoría", "Marca"], rows)

            elif sid == "pedidos":
                vista = api.get("/v1/pedidos/admin/vista", params={"limit": 500}) or []
                rows = []
                for p in vista:
                    if not isinstance(p, dict):
                        continue
                    rows.append(
                        [
                            _safe(p.get("folio")),
                            _safe(p.get("cliente_nombre")),
                            _safe(p.get("destino")),
                            _safe(p.get("piezas")),
                            _safe(p.get("estatus")),
                            _safe(p.get("estatus_modulo")),
                            _safe(p.get("total")),
                            (_safe(p.get("fecha_pedido"))[:19] if p.get("fecha_pedido") else "—"),
                        ]
                    )
                add_table(
                    "Pedidos",
                    ["Folio", "Cliente", "Destino (municipio/estado)", "Piezas", "Estatus", "Módulo", "Total MXN", "Fecha"],
                    rows,
                )

            elif sid == "parametros" and rol == "Administrador":
                pars = api.get("/v1/parametros-sistema/") or []
                rows = [
                    [_safe(p.get("tipo")), _safe(p.get("clave")), _safe(p.get("valor")), "Sí" if p.get("activo") else "No"]
                    for p in pars
                ]
                add_table("Parámetros de sistema", ["Tipo", "Clave", "Valor", "Activo"], rows)

            elif sid == "mensajes_contacto" and rol == "Administrador":
                msgs = api.get("/v1/portal-contacto/mensajes") or []
                rows = []
                for m in msgs:
                    if not isinstance(m, dict):
                        continue
                    ar = (m.get("admin_reply") or "").strip()
                    reply_snip = (ar[:200] + "…") if len(ar) > 200 else (ar or "—")
                    msg_txt = _safe(m.get("mensaje"))
                    msg_snip = (msg_txt[:120] + "…") if len(msg_txt) > 120 else msg_txt
                    estado = "Contestado" if ar else ("Leído" if m.get("is_read") else "No leído")
                    rows.append(
                        [
                            _safe(m.get("id")),
                            (_safe(m.get("creado_en"))[:19] if m.get("creado_en") else "—"),
                            _safe(m.get("nombre")),
                            _safe(m.get("email")),
                            _safe(m.get("asunto")),
                            estado,
                            msg_snip,
                            reply_snip,
                        ]
                    )
                add_table(
                    "Mensajes de contacto (portal web)",
                    ["ID", "Fecha", "Nombre", "Correo", "Asunto", "Estado", "Mensaje (extracto)", "Respuesta admin (extracto)"],
                    rows,
                )

            elif sid == "guias" and rol in ("Administrador", "Logística"):
                guias = api.get("/v1/guias-envio/") or []
                rows = []
                for g in guias:
                    if not isinstance(g, dict):
                        g = {k: getattr(g, k, None) for k in ("pedido_id", "tipo", "paqueteria", "numero_rastreo")}
                    pid = g.get("pedido_id")
                    folio = maps["pedido_folio"].get(int(pid), f"Pedido #{pid}") if pid is not None else "—"
                    rows.append([folio, _safe(g.get("tipo")), _safe(g.get("paqueteria")), _safe(g.get("numero_rastreo"))])
                add_table("Guías de envío", ["Pedido (folio)", "Tipo", "Paquetería", "Número de rastreo"], rows)

            elif sid == "direcciones" and rol in ("Administrador", "Logística"):
                dirs = api.get("/v1/direcciones/") or []
                rows = []
                for d in dirs:
                    if not isinstance(d, dict):
                        d = {
                            k: getattr(d, k, None)
                            for k in ("calle_principal", "num_ext", "colonia", "municipio", "estado", "cp", "cliente_id")
                        }
                    cid = d.get("cliente_id")
                    cli = maps["cliente"].get(int(cid), "—") if cid is not None else "—"
                    rows.append(
                        [
                            cli,
                            _safe(d.get("calle_principal")),
                            _safe(d.get("num_ext")),
                            _safe(d.get("colonia")),
                            _safe(d.get("municipio")),
                            _safe(d.get("estado")),
                            _safe(d.get("cp")),
                        ]
                    )
                add_table("Direcciones", ["Cliente", "Calle", "Núm. ext.", "Colonia", "Municipio", "Estado", "C.P."], rows)

            elif sid == "inventarios" and rol in ("Administrador", "Almacén"):
                inv = api.get("/v1/inventarios/") or []
                rows = []
                for i in inv:
                    if not isinstance(i, dict):
                        i = {k: getattr(i, k, None) for k in ("autoparte_id", "stock_actual", "stock_minimo", "ubicacion_id")}
                    aid, uid = i.get("autoparte_id"), i.get("ubicacion_id")
                    ap_lbl = maps["autoparte"].get(int(aid), f"Autoparte #{aid}") if aid is not None else "—"
                    ubi = maps["ubicacion"].get(int(uid), "—") if uid is not None else "Sin ubicación"
                    rows.append([ap_lbl, _safe(i.get("stock_actual")), _safe(i.get("stock_minimo")), ubi])
                add_table("Inventarios", ["Autoparte", "Stock actual", "Stock mínimo", "Ubicación (pasillo-estante)"], rows)

            elif sid == "ubicaciones" and rol in ("Administrador", "Almacén"):
                ubi = api.get("/v1/ubicaciones/") or []
                rows = []
                for u in ubi:
                    if not isinstance(u, dict):
                        u = {k: getattr(u, k, None) for k in ("pasillo", "estante", "nivel", "capacidad", "descripcion", "activo")}
                    rows.append(
                        [
                            _safe(u.get("pasillo")),
                            _safe(u.get("estante")),
                            _safe(u.get("nivel")),
                            _safe(u.get("capacidad")),
                            (_safe(u.get("descripcion"))[:60] + "…")
                            if u.get("descripcion") and len(_safe(u.get("descripcion"))) > 60
                            else _safe(u.get("descripcion")),
                            "Sí" if u.get("activo") else "No",
                        ]
                    )
                add_table("Ubicaciones de almacén", ["Pasillo", "Estante", "Nivel", "Capacidad", "Descripción", "Activo"], rows)

    except Exception as exc:
        sections.append({"type": "text", "title": "Error al cargar datos", "body": str(exc)})

    return {"title": title, "rol": rol, "sections": sections, "generated_at": datetime.now(timezone.utc).isoformat()}


def _pdf_footer(canvas: Any, doc: Any) -> None:
    canvas.saveState()
    canvas.setFont(_FONT_BODY, 8)
    canvas.setFillColor(colors.HexColor(MACUIN_NAVY))
    w = A4[0]
    canvas.drawString(doc.leftMargin, 0.65 * cm, "MACUIN Autopartes — documento interno")
    canvas.drawRightString(w - doc.rightMargin, 0.65 * cm, f"Página {canvas.getPageNumber()}")
    canvas.restoreState()


def render_pdf(context: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        title=context.get("title", "Reporte")[:60],
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.15 * cm,
        bottomMargin=1.85 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        name="RepTitle",
        parent=styles["Title"],
        fontName=_FONT_TITLE,
        fontSize=18,
        textColor=colors.HexColor(MACUIN_NAVY_DARK),
        spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        name="RepH2",
        parent=styles["Heading2"],
        fontName=_FONT_TITLE,
        fontSize=12,
        textColor=colors.HexColor(MACUIN_NAVY),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        name="RepBody",
        parent=styles["BodyText"],
        fontName=_FONT_BODY,
        fontSize=9,
        leading=12,
        textColor=colors.HexColor(MACUIN_SLATE),
    )
    sub_style = ParagraphStyle(
        name="RepSub",
        parent=body_style,
        fontSize=8.5,
        textColor=colors.HexColor(MACUIN_MUTED),
        spaceAfter=4,
    )

    story: List[Any] = []
    brand = Table([[Paragraph(escape(context.get("title", "Reporte")), title_style)]], colWidths=[17 * cm])
    brand.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(MACUIN_GOLD)),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor(MACUIN_GOLD_DEEP)),
            ]
        )
    )
    story.append(brand)
    story.append(Spacer(1, 8))

    for sec in context.get("sections", []):
        if sec["type"] == "text":
            story.append(Paragraph(escape(sec["title"]), h2_style))
            story.append(Paragraph(escape(sec.get("body", "")).replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 8))
        elif sec["type"] == "chart":
            story.append(Paragraph(escape(sec.get("title", "Gráfica")), h2_style))
            if sec.get("subtitle"):
                story.append(Paragraph(escape(_safe(sec["subtitle"])), sub_style))
            png = sec.get("png_bytes") or b""
            if png:
                bio = io.BytesIO(png)
                bio.seek(0)
                story.append(RLImage(bio, width=15.6 * cm, height=7.85 * cm))
            story.append(Spacer(1, 12))
        elif sec["type"] == "table":
            story.append(Paragraph(escape(sec["title"]), h2_style))
            hdr = [Paragraph(f"<b>{escape(str(h))}</b>", body_style) for h in sec["headers"]]
            body_rows: List[List[Any]] = [hdr]
            for row in sec["rows"]:
                body_rows.append([Paragraph(escape(_safe(c)), body_style) for c in row])
            t = Table(body_rows, repeatRows=1, hAlign="LEFT")
            t.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(MACUIN_NAVY)),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#f8fafc")),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("FONTNAME", (0, 0), (-1, 0), _FONT_TITLE),
                        ("FONTSIZE", (0, 0), (-1, 0), 9),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                        ("TOPPADDING", (0, 0), (-1, 0), 8),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 1), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#cbd5e1")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
                        ("LINEBELOW", (0, 0), (-1, 0), 1.2, colors.HexColor(MACUIN_NAVY_DARK)),
                    ]
                )
            )
            story.append(t)
            story.append(Spacer(1, 14))

    doc.build(story, onFirstPage=_pdf_footer, onLaterPages=_pdf_footer)
    return buf.getvalue()


def _docx_set_cell_shading(cell: Any, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex.replace("#", ""))
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_style_table_header(table: Any) -> None:
    from docx.shared import Pt, RGBColor

    hdr = table.rows[0].cells
    for cell in hdr:
        _docx_set_cell_shading(cell, MACUIN_NAVY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(248, 250, 252)


def render_docx(context: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    h0 = doc.add_heading(_safe(context.get("title", "Reporte")), level=0)
    for r in h0.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
    sub = doc.add_paragraph()
    sub.add_run("MACUIN Autopartes").bold = True
    sub.add_run(f" · Rol: {_safe(context.get('rol'))} · {_safe(context.get('generated_at'))}")
    for run in sub.runs[1:]:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(51, 65, 85)

    for sec in context.get("sections", []):
        if sec["type"] == "text":
            h = doc.add_heading(_safe(sec.get("title", "")), level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(30, 58, 95)
            for line in (sec.get("body") or "").split("\n"):
                doc.add_paragraph(line)
        elif sec["type"] == "chart":
            h = doc.add_heading(_safe(sec.get("title", "Gráfica")), level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(30, 58, 95)
            if sec.get("subtitle"):
                p = doc.add_paragraph(_safe(sec["subtitle"]))
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(100, 116, 139)
            png = sec.get("png_bytes") or b""
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(6.45))
            doc.add_paragraph()
        elif sec["type"] == "table":
            h = doc.add_heading(_safe(sec.get("title", "Tabla")), level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(30, 58, 95)
            headers = sec.get("headers") or []
            rows = sec.get("rows") or []
            if not headers:
                continue
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, htxt in enumerate(headers):
                hdr_cells[i].text = _safe(htxt)
            _docx_style_table_header(table)
            for ri, data_row in enumerate(rows, start=1):
                for ci in range(len(headers)):
                    val = data_row[ci] if ci < len(data_row) else ""
                    table.rows[ri].cells[ci].text = _safe(val)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx_sheet_name(title: str, used: Dict[str, int]) -> str:
    base = re.sub(r"[^\w\s\-]", "", title, flags=re.UNICODE) or "Hoja"
    base = re.sub(r"\s+", " ", base).strip()[:28] or "Hoja"
    name = base
    n = used.get(name, 0)
    used[name] = n + 1
    if n:
        suffix = f"_{n + 1}"
        name = (base[: 31 - len(suffix)] + suffix)[:31]
    return name


def render_xlsx(context: Dict[str, Any]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as SheetImage
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    navy_fill = PatternFill(start_color="1E3A5F", end_color="1E3A5F", fill_type="solid")
    gold_fill = PatternFill(start_color="FBBF24", end_color="FBBF24", fill_type="solid")
    hdr_font = Font(bold=True, color="F8FAFC", size=10)
    title_font = Font(bold=True, size=16, color="0F172A")
    brand_font = Font(bold=True, size=11, color="0F172A")

    wb = Workbook()
    used_titles: Dict[str, int] = {}
    meta = wb.active
    meta.title = _xlsx_sheet_name("Portada", used_titles)
    meta.merge_cells("A1:F1")
    tcell = meta["A1"]
    tcell.value = _safe(context.get("title", "Reporte"))
    tcell.font = title_font
    tcell.alignment = Alignment(vertical="center", wrap_text=True)
    for col in range(1, 7):
        meta.cell(row=1, column=col).fill = gold_fill
    meta.row_dimensions[1].height = 28
    meta["A2"] = "MACUIN Autopartes — reporte interno"
    meta["A2"].font = brand_font
    meta["A3"] = f"Rol: {_safe(context.get('rol'))}"
    meta["A4"] = _safe(context.get("generated_at"))
    r = 6

    for sec in context.get("sections", []):
        if sec["type"] == "text":
            c = meta.cell(row=r, column=1, value=_safe(sec.get("title", "")))
            c.font = Font(bold=True, size=11, color="1E3A5F")
            r += 1
            for line in (sec.get("body") or "").split("\n"):
                meta.cell(row=r, column=1, value=line)
                r += 1
            r += 1
        elif sec["type"] == "chart":
            meta.cell(row=r, column=1, value=_safe(sec.get("title", "Gráfica"))).font = Font(
                bold=True, size=11, color="1E3A5F"
            )
            r += 1
            if sec.get("subtitle"):
                meta.cell(row=r, column=1, value=_safe(sec["subtitle"]))
                r += 1
            png = sec.get("png_bytes") or b""
            if png:
                img = SheetImage(io.BytesIO(png))
                img.anchor = f"B{r}"
                meta.add_image(img)
                r += 20
            r += 1
        elif sec["type"] == "table":
            stitle = _safe(sec.get("title", "Tabla"))
            ws = wb.create_sheet(title=_xlsx_sheet_name(stitle, used_titles))
            headers = sec.get("headers") or []
            for col, h in enumerate(headers, 1):
                c = ws.cell(row=1, column=col, value=_safe(h))
                c.font = hdr_font
                c.fill = navy_fill
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            for ri, data_row in enumerate(sec.get("rows") or [], start=2):
                for ci, val in enumerate(data_row, 1):
                    if ci <= len(headers):
                        ws.cell(row=ri, column=ci, value=val)
            for col in range(1, len(headers) + 1):
                letter = get_column_letter(col)
                ws.column_dimensions[letter].width = min(42, max(10, len(str(headers[col - 1])) + 4))

    meta.cell(row=r + 1, column=1, value="MACUIN Autopartes — documento interno").font = Font(
        size=9, italic=True, color="64748B"
    )

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
